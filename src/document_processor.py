"""
Document Processor for Multi-Format File Handling and Intelligent Chunking.

This module provides comprehensive document processing capabilities:
- Multi-format file support (PDF, TXT, MD, DOCX, HTML)
- Intelligent chunking with sentence boundary detection
- Comprehensive metadata extraction and preservation
- Text preprocessing and normalization
- Async batch processing with concurrency control
- Progress tracking and caching mechanisms
- Content validation and deduplication

Features:
- DocumentLoader: Abstract base for format-specific loaders
- TextChunker: Intelligent text chunking with sentence awareness
- MetadataExtractor: Comprehensive metadata extraction
- DocumentProcessor: Main orchestrator with async processing
- ProcessingCache: Efficient caching to avoid reprocessing
"""

import asyncio
import hashlib
import json
import logging
import pickle
import re
import tempfile
import time
import unicodedata
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Protocol, Tuple, Union

# File format libraries (with graceful degradation)
try:
    import pypdf
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    PyPDF2 = None
    pypdf = None

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None

try:
    from bs4 import BeautifulSoup

    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False
    BeautifulSoup = None

try:
    import chardet

    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False
    chardet = None

try:
    import aiofiles
    import aiosqlite

    ASYNC_IO_AVAILABLE = True
except ImportError:
    ASYNC_IO_AVAILABLE = False
    aiofiles = None
    aiosqlite = None

# Local imports
try:
    from vector_store import Document
except ImportError:
    # Fallback Document definition
    @dataclass
    class Document:
        page_content: str
        metadata: Dict[str, Any] = field(default_factory=dict)
        id: Optional[str] = None
        content_hash: Optional[str] = None


# Type definitions
class ProgressCallback(Protocol):
    def __call__(self, current: int, total: int, status: str = "") -> None: ...


@dataclass
class ProcessedDocument:
    """Container for processed document with metadata."""

    content: str
    metadata: Dict[str, Any]
    source_file: Path
    chunk_index: Optional[int] = None
    processing_time: float = 0.0


@dataclass
class ProcessingStats:
    """Statistics for document processing operations."""

    total_files: int
    processed_files: int
    failed_files: int
    total_chunks: int
    execution_time: float
    bytes_processed: int
    cache_hits: int = 0
    cache_misses: int = 0
    files_by_type: Dict[str, int] = field(default_factory=dict)


class DocumentProcessingError(Exception):
    """Base exception for document processing errors."""

    pass


class UnsupportedFormatError(DocumentProcessingError):
    """Raised when file format is not supported."""

    pass


class ContentValidationError(DocumentProcessingError):
    """Raised when content validation fails."""

    pass


# Abstract base classes
class DocumentLoader(ABC):
    """Abstract base class for document loaders."""

    @abstractmethod
    def can_load(self, file_path: Path) -> bool:
        """Check if this loader can handle the file."""
        pass

    @abstractmethod
    async def load(self, file_path: Path, **kwargs) -> Tuple[str, Dict[str, Any]]:
        """Load document content and metadata."""
        pass

    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        pass


class TextPreprocessor:
    """Handles text preprocessing and normalization."""

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.TextPreprocessor")

        # Compile regex patterns for efficiency
        self.control_char_pattern = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]")
        self.whitespace_pattern = re.compile(r"\s+")
        self.paragraph_pattern = re.compile(r"\n\s*\n")

    def normalize_text(self, text: str) -> str:
        """
        Normalize text with comprehensive preprocessing.

        Args:
            text: Raw text content

        Returns:
            Normalized text
        """
        if not text:
            return ""

        try:
            # Unicode normalization
            text = unicodedata.normalize("NFKC", text)

            # Remove control characters except newlines and tabs
            text = self.control_char_pattern.sub("", text)

            # Normalize whitespace while preserving paragraph structure
            paragraphs = self.paragraph_pattern.split(text)
            normalized_paragraphs = []

            for paragraph in paragraphs:
                # Normalize whitespace within paragraph
                normalized = self.whitespace_pattern.sub(" ", paragraph.strip())
                if normalized:  # Only add non-empty paragraphs
                    normalized_paragraphs.append(normalized)

            # Rejoin paragraphs with double newlines
            text = "\n\n".join(normalized_paragraphs)

            return text.strip()

        except Exception as e:
            self.logger.warning(f"Text normalization failed: {e}")
            return text.strip()

    def detect_encoding(self, file_path: Path) -> str:
        """
        Detect file encoding with fallback strategies.

        Args:
            file_path: Path to file

        Returns:
            Detected encoding
        """
        if CHARDET_AVAILABLE:
            try:
                with open(file_path, "rb") as f:
                    raw_data = f.read(10000)  # Read first 10KB for detection
                    result = chardet.detect(raw_data)
                    if result["encoding"] and result["confidence"] > 0.7:
                        return result["encoding"]
            except Exception as e:
                self.logger.warning(f"Character detection failed for {file_path}: {e}")

        # Fallback encodings to try
        fallback_encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "ascii"]

        for encoding in fallback_encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    f.read(1000)  # Try to read a small portion
                return encoding
            except (UnicodeDecodeError, UnicodeError):
                continue

        self.logger.warning(
            f"Could not detect encoding for {file_path}, using utf-8 with errors='replace'"
        )
        return "utf-8"


class TextChunker:
    """Intelligent text chunking with sentence boundary detection."""

    def __init__(
        self, chunk_size: int = 1000, overlap: int = 200, min_chunk_size: int = 100
    ):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.min_chunk_size = min_chunk_size
        self.logger = logging.getLogger(f"{__name__}.TextChunker")

        # Sentence boundary patterns
        self.sentence_pattern = re.compile(
            r'(?<=[.!?])\s+(?=[A-Z])|(?<=[.!?]["\']\s)\s*(?=[A-Z])', re.MULTILINE
        )

    def chunk_text(self, text: str, preserve_paragraphs: bool = True) -> List[str]:
        """
        Chunk text with intelligent sentence boundary detection.

        Args:
            text: Input text to chunk
            preserve_paragraphs: Whether to respect paragraph boundaries

        Returns:
            List of text chunks
        """
        if not text or len(text) <= self.chunk_size:
            return [text] if text.strip() else []

        try:
            if preserve_paragraphs:
                return self._chunk_with_paragraphs(text)
            else:
                return self._chunk_with_sentences(text)

        except Exception as e:
            self.logger.warning(
                f"Chunking failed, falling back to simple splitting: {e}"
            )
            return self._simple_chunk(text)

    def _chunk_with_paragraphs(self, text: str) -> List[str]:
        """Chunk text while preserving paragraph boundaries."""
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = ""

        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            # If paragraph alone exceeds chunk size, split it
            if len(paragraph) > self.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""

                # Split long paragraph by sentences
                para_chunks = self._chunk_with_sentences(paragraph)
                chunks.extend(para_chunks)
                continue

            # Check if adding paragraph would exceed chunk size
            potential_chunk = f"{current_chunk}\n\n{paragraph}".strip()

            if len(potential_chunk) <= self.chunk_size:
                current_chunk = potential_chunk
            else:
                # Save current chunk and start new one
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph

        # Add final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return self._apply_overlap(chunks)

    def _chunk_with_sentences(self, text: str) -> List[str]:
        """Chunk text by sentence boundaries."""
        sentences = self.sentence_pattern.split(text)
        if len(sentences) <= 1:
            return self._simple_chunk(text)

        chunks = []
        current_chunk = ""

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            potential_chunk = f"{current_chunk} {sentence}".strip()

            if len(potential_chunk) <= self.chunk_size:
                current_chunk = potential_chunk
            else:
                # If single sentence is too long, split it
                if len(sentence) > self.chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk)
                    chunks.extend(self._simple_chunk(sentence))
                    current_chunk = ""
                else:
                    # Save current chunk and start new one
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = sentence

        # Add final chunk
        if current_chunk:
            chunks.append(current_chunk)

        return self._apply_overlap(chunks)

    def _simple_chunk(self, text: str) -> List[str]:
        """Simple character-based chunking as fallback."""
        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end].strip()

            if chunk:
                chunks.append(chunk)

            start = end - self.overlap if self.overlap > 0 else end

        return chunks

    def _apply_overlap(self, chunks: List[str]) -> List[str]:
        """Apply overlap between chunks."""
        if not chunks or self.overlap <= 0:
            return chunks

        overlapped_chunks = [chunks[0]]

        for i in range(1, len(chunks)):
            prev_chunk = chunks[i - 1]
            current_chunk = chunks[i]

            # Get overlap text from previous chunk
            overlap_text = (
                prev_chunk[-self.overlap :]
                if len(prev_chunk) > self.overlap
                else prev_chunk
            )

            # Combine with current chunk
            combined = f"{overlap_text} {current_chunk}".strip()
            overlapped_chunks.append(combined)

        return overlapped_chunks


class MetadataExtractor:
    """Extracts comprehensive metadata from files."""

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.MetadataExtractor")

    def extract_file_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract comprehensive file metadata.

        Args:
            file_path: Path to file

        Returns:
            Dictionary of metadata
        """
        try:
            stat = file_path.stat()

            metadata = {
                "source": str(file_path.absolute()),
                "filename": file_path.name,
                "file_extension": file_path.suffix.lower(),
                "file_size": stat.st_size,
                "creation_time": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modification_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "processing_time": datetime.utcnow().isoformat(),
                "file_type": self._determine_file_type(file_path),
            }

            # Add file hash for deduplication
            metadata["file_hash"] = self._calculate_file_hash(file_path)

            return metadata

        except Exception as e:
            self.logger.warning(f"Failed to extract metadata for {file_path}: {e}")
            return {
                "source": str(file_path.absolute()),
                "filename": file_path.name,
                "file_extension": file_path.suffix.lower(),
                "error": str(e),
            }

    def _determine_file_type(self, file_path: Path) -> str:
        """Determine file type from extension."""
        extension = file_path.suffix.lower()

        type_mapping = {
            ".pdf": "pdf",
            ".txt": "text",
            ".md": "markdown",
            ".docx": "docx",
            ".doc": "doc",
            ".html": "html",
            ".htm": "html",
            ".xml": "xml",
            ".rtf": "rtf",
        }

        return type_mapping.get(extension, "unknown")

    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file content."""
        try:
            hash_sha256 = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception as e:
            self.logger.warning(f"Failed to calculate hash for {file_path}: {e}")
            return ""


# Specific document loaders
class PDFLoader(DocumentLoader):
    """PDF document loader with multiple parsing strategies."""

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.PDFLoader")

    def can_load(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf" and PDF_AVAILABLE

    def get_supported_extensions(self) -> List[str]:
        return [".pdf"] if PDF_AVAILABLE else []

    async def load(self, file_path: Path, **kwargs) -> Tuple[str, Dict[str, Any]]:
        """Load PDF content with metadata."""
        if not PDF_AVAILABLE:
            raise UnsupportedFormatError("PDF libraries not available")

        try:
            # Try pypdf first (newer, more reliable)
            if pypdf:
                return await self._load_with_pypdf(file_path)
            elif PyPDF2:
                return await self._load_with_pypdf2(file_path)
            else:
                raise UnsupportedFormatError("No PDF library available")

        except Exception as e:
            self.logger.error(f"Failed to load PDF {file_path}: {e}")
            raise DocumentProcessingError(f"PDF loading failed: {e}")

    async def _load_with_pypdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Load PDF using pypdf library."""
        import pypdf

        text_content = []
        metadata = {"pages": []}

        try:
            with open(file_path, "rb") as file:
                pdf_reader = pypdf.PdfReader(file)

                # Extract document metadata
                if pdf_reader.metadata:
                    metadata.update(
                        {
                            "title": getattr(pdf_reader.metadata, "title", ""),
                            "author": getattr(pdf_reader.metadata, "author", ""),
                            "subject": getattr(pdf_reader.metadata, "subject", ""),
                            "creator": getattr(pdf_reader.metadata, "creator", ""),
                            "producer": getattr(pdf_reader.metadata, "producer", ""),
                            "creation_date": getattr(
                                pdf_reader.metadata, "creation_date", ""
                            ),
                            "modification_date": getattr(
                                pdf_reader.metadata, "modification_date", ""
                            ),
                        }
                    )

                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                            metadata["pages"].append(
                                {
                                    "page_number": page_num + 1,
                                    "char_count": len(page_text),
                                }
                            )
                    except Exception as e:
                        self.logger.warning(
                            f"Failed to extract text from page {page_num + 1}: {e}"
                        )

                metadata["total_pages"] = len(pdf_reader.pages)

        except Exception as e:
            raise DocumentProcessingError(f"pypdf extraction failed: {e}")

        combined_text = "\n\n".join(text_content)
        return combined_text, metadata

    async def _load_with_pypdf2(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Load PDF using PyPDF2 library (fallback)."""
        import PyPDF2

        text_content = []
        metadata = {"pages": []}

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract document metadata
                if pdf_reader.metadata:
                    metadata.update(
                        {
                            "title": pdf_reader.metadata.get("/Title", ""),
                            "author": pdf_reader.metadata.get("/Author", ""),
                            "subject": pdf_reader.metadata.get("/Subject", ""),
                            "creator": pdf_reader.metadata.get("/Creator", ""),
                            "producer": pdf_reader.metadata.get("/Producer", ""),
                            "creation_date": pdf_reader.metadata.get(
                                "/CreationDate", ""
                            ),
                            "modification_date": pdf_reader.metadata.get(
                                "/ModDate", ""
                            ),
                        }
                    )

                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                            metadata["pages"].append(
                                {
                                    "page_number": page_num + 1,
                                    "char_count": len(page_text),
                                }
                            )
                    except Exception as e:
                        self.logger.warning(
                            f"Failed to extract text from page {page_num + 1}: {e}"
                        )

                metadata["total_pages"] = len(pdf_reader.pages)

        except Exception as e:
            raise DocumentProcessingError(f"PyPDF2 extraction failed: {e}")

        combined_text = "\n\n".join(text_content)
        return combined_text, metadata


class TextLoader(DocumentLoader):
    """Text file loader with encoding detection."""

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.TextLoader")
        self.preprocessor = TextPreprocessor()

    def can_load(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".txt", ".md"]

    def get_supported_extensions(self) -> List[str]:
        return [".txt", ".md"]

    async def load(self, file_path: Path, **kwargs) -> Tuple[str, Dict[str, Any]]:
        """Load text file with encoding detection."""
        encoding = self.preprocessor.detect_encoding(file_path)

        try:
            if ASYNC_IO_AVAILABLE:
                async with aiofiles.open(
                    file_path, "r", encoding=encoding, errors="replace"
                ) as f:
                    content = await f.read()
            else:
                with open(file_path, "r", encoding=encoding, errors="replace") as f:
                    content = f.read()

            metadata = {
                "encoding": encoding,
                "line_count": content.count("\n") + 1,
                "char_count": len(content),
            }

            return content, metadata

        except Exception as e:
            self.logger.error(f"Failed to load text file {file_path}: {e}")
            raise DocumentProcessingError(f"Text file loading failed: {e}")


class DocxLoader(DocumentLoader):
    """DOCX document loader."""

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.DocxLoader")

    def can_load(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx" and DOCX_AVAILABLE

    def get_supported_extensions(self) -> List[str]:
        return [".docx"] if DOCX_AVAILABLE else []

    async def load(self, file_path: Path, **kwargs) -> Tuple[str, Dict[str, Any]]:
        """Load DOCX content with metadata."""
        if not DOCX_AVAILABLE:
            raise UnsupportedFormatError("python-docx library not available")

        try:
            doc = DocxDocument(file_path)

            # Extract text content
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            content = "\n\n".join(paragraphs)

            # Extract metadata
            metadata = {"paragraph_count": len(paragraphs), "char_count": len(content)}

            # Document properties
            if doc.core_properties:
                props = doc.core_properties
                metadata.update(
                    {
                        "title": props.title or "",
                        "author": props.author or "",
                        "subject": props.subject or "",
                        "keywords": props.keywords or "",
                        "comments": props.comments or "",
                        "created": props.created.isoformat() if props.created else "",
                        "modified": (
                            props.modified.isoformat() if props.modified else ""
                        ),
                        "last_modified_by": props.last_modified_by or "",
                    }
                )

            return content, metadata

        except Exception as e:
            self.logger.error(f"Failed to load DOCX {file_path}: {e}")
            raise DocumentProcessingError(f"DOCX loading failed: {e}")


class HTMLLoader(DocumentLoader):
    """HTML document loader with content extraction."""

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.HTMLLoader")
        self.preprocessor = TextPreprocessor()

    def can_load(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".html", ".htm"] and HTML_AVAILABLE

    def get_supported_extensions(self) -> List[str]:
        return [".html", ".htm"] if HTML_AVAILABLE else []

    async def load(self, file_path: Path, **kwargs) -> Tuple[str, Dict[str, Any]]:
        """Load HTML content and extract text."""
        if not HTML_AVAILABLE:
            raise UnsupportedFormatError("BeautifulSoup library not available")

        try:
            encoding = self.preprocessor.detect_encoding(file_path)

            if ASYNC_IO_AVAILABLE:
                async with aiofiles.open(
                    file_path, "r", encoding=encoding, errors="replace"
                ) as f:
                    html_content = await f.read()
            else:
                with open(file_path, "r", encoding=encoding, errors="replace") as f:
                    html_content = f.read()

            soup = BeautifulSoup(html_content, "html.parser")

            # Extract metadata from HTML head
            metadata = {"encoding": encoding, "html_length": len(html_content)}

            # Extract title
            title_tag = soup.find("title")
            if title_tag:
                metadata["title"] = title_tag.get_text().strip()

            # Extract meta tags
            meta_tags = soup.find_all("meta")
            for meta in meta_tags:
                name = meta.get("name") or meta.get("property")
                content = meta.get("content")
                if name and content:
                    metadata[f"meta_{name}"] = content

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Extract text content
            text = soup.get_text()

            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)

            metadata["extracted_char_count"] = len(text)

            return text, metadata

        except Exception as e:
            self.logger.error(f"Failed to load HTML {file_path}: {e}")
            raise DocumentProcessingError(f"HTML loading failed: {e}")


class ProcessingCache:
    """Efficient caching system to avoid reprocessing identical files."""

    def __init__(self, cache_dir: Optional[Path] = None):
        self.cache_dir = (
            cache_dir or Path(tempfile.gettempdir()) / "document_processing_cache"
        )
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        self.cache_db_path = self.cache_dir / "processing_cache.db"
        self._initialized = False
        self.logger = logging.getLogger(f"{__name__}.ProcessingCache")

    async def initialize(self) -> None:
        """Initialize cache database."""
        if self._initialized:
            return

        if ASYNC_IO_AVAILABLE:
            async with aiosqlite.connect(self.cache_db_path) as db:
                await db.execute(
                    """
                    CREATE TABLE IF NOT EXISTS file_cache (
                        file_hash TEXT PRIMARY KEY,
                        file_path TEXT NOT NULL,
                        file_size INTEGER NOT NULL,
                        modification_time TEXT NOT NULL,
                        processed_chunks INTEGER NOT NULL,
                        processing_time REAL NOT NULL,
                        cache_data BLOB NOT NULL,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        access_count INTEGER DEFAULT 1,
                        last_accessed TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )
                """
                )
                await db.execute(
                    """
                    CREATE INDEX IF NOT EXISTS idx_file_path ON file_cache(file_path)
                """
                )
                await db.execute(
                    """
                    CREATE INDEX IF NOT EXISTS idx_modification_time ON file_cache(modification_time)
                """
                )
                await db.commit()
        else:
            # Fallback to synchronous operations
            import sqlite3

            with sqlite3.connect(self.cache_db_path) as db:
                db.execute(
                    """
                    CREATE TABLE IF NOT EXISTS file_cache (
                        file_hash TEXT PRIMARY KEY,
                        file_path TEXT NOT NULL,
                        file_size INTEGER NOT NULL,
                        modification_time TEXT NOT NULL,
                        processed_chunks INTEGER NOT NULL,
                        processing_time REAL NOT NULL,
                        cache_data BLOB NOT NULL,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        access_count INTEGER DEFAULT 1,
                        last_accessed TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )
                """
                )
                db.execute(
                    """
                    CREATE INDEX IF NOT EXISTS idx_file_path ON file_cache(file_path)
                """
                )
                db.execute(
                    """
                    CREATE INDEX IF NOT EXISTS idx_modification_time ON file_cache(modification_time)
                """
                )
                db.commit()

        self._initialized = True
        self.logger.info(f"Processing cache initialized at {self.cache_db_path}")

    async def get_cached_result(
        self, file_path: Path, file_hash: str, modification_time: str
    ) -> Optional[List[ProcessedDocument]]:
        """Retrieve cached processing result if valid."""
        if not self._initialized:
            await self.initialize()

        try:
            if ASYNC_IO_AVAILABLE:
                async with aiosqlite.connect(self.cache_db_path) as db:
                    cursor = await db.execute(
                        "SELECT cache_data FROM file_cache WHERE file_hash = ? AND modification_time = ?",
                        (file_hash, modification_time),
                    )
                    row = await cursor.fetchone()

                    if row:
                        # Update access statistics
                        await db.execute(
                            "UPDATE file_cache SET access_count = access_count + 1, last_accessed = CURRENT_TIMESTAMP WHERE file_hash = ?",
                            (file_hash,),
                        )
                        await db.commit()

                        # Deserialize cached data
                        cached_data = pickle.loads(row[0])
                        self.logger.debug(f"Cache hit for {file_path}")
                        return cached_data
            else:
                import sqlite3

                with sqlite3.connect(self.cache_db_path) as db:
                    cursor = db.execute(
                        "SELECT cache_data FROM file_cache WHERE file_hash = ? AND modification_time = ?",
                        (file_hash, modification_time),
                    )
                    row = cursor.fetchone()

                    if row:
                        db.execute(
                            "UPDATE file_cache SET access_count = access_count + 1, last_accessed = CURRENT_TIMESTAMP WHERE file_hash = ?",
                            (file_hash,),
                        )
                        db.commit()

                        cached_data = pickle.loads(row[0])
                        self.logger.debug(f"Cache hit for {file_path}")
                        return cached_data

        except Exception as e:
            self.logger.warning(f"Cache retrieval failed for {file_path}: {e}")

        return None

    async def store_result(
        self,
        file_path: Path,
        file_hash: str,
        file_size: int,
        modification_time: str,
        processed_documents: List[ProcessedDocument],
        processing_time: float,
    ) -> None:
        """Store processing result in cache."""
        if not self._initialized:
            await self.initialize()

        try:
            cache_data = pickle.dumps(processed_documents)

            if ASYNC_IO_AVAILABLE:
                async with aiosqlite.connect(self.cache_db_path) as db:
                    await db.execute(
                        "INSERT OR REPLACE INTO file_cache (file_hash, file_path, file_size, modification_time, processed_chunks, processing_time, cache_data) VALUES (?, ?, ?, ?, ?, ?, ?)",
                        (
                            file_hash,
                            str(file_path),
                            file_size,
                            modification_time,
                            len(processed_documents),
                            processing_time,
                            cache_data,
                        ),
                    )
                    await db.commit()
            else:
                import sqlite3

                with sqlite3.connect(self.cache_db_path) as db:
                    db.execute(
                        "INSERT OR REPLACE INTO file_cache (file_hash, file_path, file_size, modification_time, processed_chunks, processing_time, cache_data) VALUES (?, ?, ?, ?, ?, ?, ?)",
                        (
                            file_hash,
                            str(file_path),
                            file_size,
                            modification_time,
                            len(processed_documents),
                            processing_time,
                            cache_data,
                        ),
                    )
                    db.commit()

            self.logger.debug(f"Cached result for {file_path}")

        except Exception as e:
            self.logger.warning(f"Cache storage failed for {file_path}: {e}")

    async def get_cache_stats(self) -> Dict[str, Any]:
        """Get cache statistics."""
        if not self._initialized:
            await self.initialize()

        try:
            if ASYNC_IO_AVAILABLE:
                async with aiosqlite.connect(self.cache_db_path) as db:
                    cursor = await db.execute(
                        "SELECT COUNT(*), SUM(file_size), SUM(processed_chunks), AVG(processing_time), SUM(access_count) FROM file_cache"
                    )
                    stats = await cursor.fetchone()

                    cursor = await db.execute(
                        "SELECT COUNT(*) FROM file_cache WHERE created_at > datetime('now', '-1 day')"
                    )
                    recent_count = (await cursor.fetchone())[0]
            else:
                import sqlite3

                with sqlite3.connect(self.cache_db_path) as db:
                    cursor = db.execute(
                        "SELECT COUNT(*), SUM(file_size), SUM(processed_chunks), AVG(processing_time), SUM(access_count) FROM file_cache"
                    )
                    stats = cursor.fetchone()

                    cursor = db.execute(
                        "SELECT COUNT(*) FROM file_cache WHERE created_at > datetime('now', '-1 day')"
                    )
                    recent_count = cursor.fetchone()[0]

            return {
                "cached_files": stats[0] or 0,
                "total_file_size": stats[1] or 0,
                "total_chunks": stats[2] or 0,
                "avg_processing_time": round(stats[3] or 0, 3),
                "total_access_count": stats[4] or 0,
                "recent_additions": recent_count or 0,
                "cache_size_mb": (
                    self.cache_db_path.stat().st_size / (1024 * 1024)
                    if self.cache_db_path.exists()
                    else 0
                ),
            }

        except Exception as e:
            self.logger.warning(f"Failed to get cache stats: {e}")
            return {}

    async def clear_cache(self, older_than_days: Optional[int] = None) -> int:
        """Clear cache entries."""
        if not self._initialized:
            await self.initialize()

        try:
            if ASYNC_IO_AVAILABLE:
                async with aiosqlite.connect(self.cache_db_path) as db:
                    if older_than_days:
                        cursor = await db.execute(
                            "DELETE FROM file_cache WHERE created_at < datetime('now', ? || ' days')",
                            (f"-{older_than_days}",),
                        )
                    else:
                        cursor = await db.execute("DELETE FROM file_cache")

                    deleted_count = cursor.rowcount
                    await db.commit()
            else:
                import sqlite3

                with sqlite3.connect(self.cache_db_path) as db:
                    if older_than_days:
                        cursor = db.execute(
                            "DELETE FROM file_cache WHERE created_at < datetime('now', ? || ' days')",
                            (f"-{older_than_days}",),
                        )
                    else:
                        cursor = db.execute("DELETE FROM file_cache")

                    deleted_count = cursor.rowcount
                    db.commit()

            return deleted_count

        except Exception as e:
            self.logger.error(f"Cache clearing failed: {e}")
            return 0


class DocumentProcessor:
    """
    Main document processor with multi-format support and intelligent chunking.

    Features:
    - Multi-format file support (PDF, TXT, MD, DOCX, HTML)
    - Intelligent chunking with sentence boundaries
    - Comprehensive metadata extraction
    - Async batch processing with concurrency control
    - Progress tracking and caching
    - Content validation and deduplication
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        overlap: int = 200,
        max_concurrency: int = 5,
        cache_dir: Optional[Path] = None,
        custom_tags: Optional[Dict[str, str]] = None,
    ):
        """
        Initialize DocumentProcessor.

        Args:
            chunk_size: Target size for text chunks
            overlap: Overlap between chunks
            max_concurrency: Maximum concurrent file processing
            cache_dir: Directory for processing cache
            custom_tags: Custom tags to add to all documents
        """
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.max_concurrency = max_concurrency
        self.custom_tags = custom_tags or {}

        # Initialize components
        self.text_chunker = TextChunker(chunk_size, overlap)
        self.metadata_extractor = MetadataExtractor()
        self.text_preprocessor = TextPreprocessor()
        self.cache = ProcessingCache(cache_dir)

        # Initialize document loaders
        self.loaders = [PDFLoader(), TextLoader(), DocxLoader(), HTMLLoader()]

        # Semaphore for concurrency control
        self.semaphore = asyncio.Semaphore(max_concurrency)

        self.logger = logging.getLogger(f"{__name__}.DocumentProcessor")

    def get_supported_formats(self) -> List[str]:
        """Get list of all supported file formats."""
        formats = []
        for loader in self.loaders:
            formats.extend(loader.get_supported_extensions())
        return sorted(set(formats))

    def can_process(self, file_path: Path) -> bool:
        """Check if file can be processed."""
        return any(loader.can_load(file_path) for loader in self.loaders)

    def _get_loader(self, file_path: Path) -> Optional[DocumentLoader]:
        """Get appropriate loader for file."""
        for loader in self.loaders:
            if loader.can_load(file_path):
                return loader
        return None

    def _validate_content(self, content: str) -> bool:
        """Validate document content."""
        if not content or not content.strip():
            return False

        # Check minimum content length
        if len(content.strip()) < 10:
            return False

        # Check for reasonable text content (not mostly special characters)
        text_chars = sum(c.isalnum() or c.isspace() for c in content)
        if text_chars / len(content) < 0.5:
            return False

        return True

    def _clean_content(self, content: str) -> str:
        """Clean and normalize content."""
        # Normalize text
        content = self.text_preprocessor.normalize_text(content)

        # Remove excessive whitespace
        lines = content.split("\n")
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            if line:  # Only keep non-empty lines
                cleaned_lines.append(line)

        return "\n".join(cleaned_lines)

    async def process_file(
        self,
        file_path: Path,
        preserve_paragraphs: bool = True,
        custom_metadata: Optional[Dict[str, Any]] = None,
    ) -> List[ProcessedDocument]:
        """
        Process single file into chunks with metadata.

        Args:
            file_path: Path to file to process
            preserve_paragraphs: Whether to preserve paragraph boundaries
            custom_metadata: Additional metadata to include

        Returns:
            List of processed document chunks
        """
        async with self.semaphore:
            start_time = time.time()

            try:
                # Check if file can be processed
                if not self.can_process(file_path):
                    raise UnsupportedFormatError(
                        f"Unsupported file format: {file_path.suffix}"
                    )

                # Extract file metadata
                file_metadata = self.metadata_extractor.extract_file_metadata(file_path)

                # Check cache first
                cached_result = await self.cache.get_cached_result(
                    file_path,
                    file_metadata.get("file_hash", ""),
                    file_metadata.get("modification_time", ""),
                )

                if cached_result:
                    self.logger.debug(f"Using cached result for {file_path}")
                    return cached_result

                # Get appropriate loader
                loader = self._get_loader(file_path)
                if not loader:
                    raise UnsupportedFormatError(f"No loader available for {file_path}")

                # Load document content
                self.logger.debug(f"Loading content from {file_path}")
                content, document_metadata = await loader.load(file_path)

                # Validate content
                if not self._validate_content(content):
                    raise ContentValidationError(
                        f"Invalid or empty content in {file_path}"
                    )

                # Clean content
                content = self._clean_content(content)

                # Chunk the content
                chunks = self.text_chunker.chunk_text(content, preserve_paragraphs)

                if not chunks:
                    raise ContentValidationError(
                        f"No valid chunks produced from {file_path}"
                    )

                # Create processed documents
                processed_docs = []

                for i, chunk in enumerate(chunks):
                    if not chunk.strip():
                        continue

                    # Combine all metadata
                    combined_metadata = {
                        **file_metadata,
                        **document_metadata,
                        **self.custom_tags,
                        **(custom_metadata or {}),
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "chunk_size": len(chunk),
                        "chunk_hash": hashlib.sha256(chunk.encode("utf-8")).hexdigest(),
                    }

                    processed_doc = ProcessedDocument(
                        content=chunk,
                        metadata=combined_metadata,
                        source_file=file_path,
                        chunk_index=i,
                        processing_time=time.time() - start_time,
                    )

                    processed_docs.append(processed_doc)

                # Cache the result
                await self.cache.store_result(
                    file_path,
                    file_metadata.get("file_hash", ""),
                    file_metadata.get("file_size", 0),
                    file_metadata.get("modification_time", ""),
                    processed_docs,
                    time.time() - start_time,
                )

                self.logger.info(
                    f"Processed {file_path}: {len(processed_docs)} chunks in {time.time() - start_time:.2f}s"
                )
                return processed_docs

            except Exception as e:
                self.logger.error(f"Failed to process {file_path}: {e}")
                raise DocumentProcessingError(f"Processing failed for {file_path}: {e}")

    async def process_files(
        self,
        file_paths: List[Path],
        preserve_paragraphs: bool = True,
        progress_callback: Optional[ProgressCallback] = None,
        custom_metadata: Optional[Dict[str, Any]] = None,
    ) -> Tuple[List[ProcessedDocument], ProcessingStats]:
        """
        Process multiple files with batch processing and progress tracking.

        Args:
            file_paths: List of file paths to process
            preserve_paragraphs: Whether to preserve paragraph boundaries
            progress_callback: Optional progress callback
            custom_metadata: Additional metadata for all files

        Returns:
            Tuple of (processed documents, processing statistics)
        """
        start_time = time.time()
        all_processed_docs = []
        stats = ProcessingStats(
            total_files=len(file_paths),
            processed_files=0,
            failed_files=0,
            total_chunks=0,
            execution_time=0.0,
            bytes_processed=0,
        )

        # Filter supported files
        supported_files = [f for f in file_paths if self.can_process(f)]
        unsupported_files = [f for f in file_paths if not self.can_process(f)]

        if unsupported_files:
            self.logger.warning(f"Skipping {len(unsupported_files)} unsupported files")

        # Count files by type
        for file_path in supported_files:
            file_type = file_path.suffix.lower()
            stats.files_by_type[file_type] = stats.files_by_type.get(file_type, 0) + 1

        # Process files concurrently
        semaphore = asyncio.Semaphore(self.max_concurrency)

        async def process_single_file(
            file_path: Path, index: int
        ) -> Optional[List[ProcessedDocument]]:
            async with semaphore:
                try:
                    processed_docs = await self.process_file(
                        file_path, preserve_paragraphs, custom_metadata
                    )

                    # Update statistics
                    stats.processed_files += 1
                    stats.total_chunks += len(processed_docs)

                    try:
                        stats.bytes_processed += file_path.stat().st_size
                    except OSError:
                        pass

                    # Call progress callback
                    if progress_callback:
                        try:
                            progress_callback(
                                stats.processed_files + stats.failed_files,
                                len(supported_files),
                                f"Processed {file_path.name}",
                            )
                        except Exception as e:
                            self.logger.warning(f"Progress callback failed: {e}")

                    return processed_docs

                except Exception as e:
                    stats.failed_files += 1
                    self.logger.error(f"Failed to process {file_path}: {e}")

                    if progress_callback:
                        try:
                            progress_callback(
                                stats.processed_files + stats.failed_files,
                                len(supported_files),
                                f"Failed: {file_path.name}",
                            )
                        except Exception as e:
                            self.logger.warning(f"Progress callback failed: {e}")

                    return None

        # Execute all processing tasks
        tasks = [
            process_single_file(file_path, i)
            for i, file_path in enumerate(supported_files)
        ]

        results = await asyncio.gather(*tasks, return_exceptions=True)

        # Collect successful results
        for result in results:
            if isinstance(result, list) and result:
                all_processed_docs.extend(result)
            elif isinstance(result, Exception):
                self.logger.error(f"Processing task failed: {result}")

        # Deduplicate documents by content hash
        unique_docs = {}
        duplicates_removed = 0

        for doc in all_processed_docs:
            content_hash = doc.metadata.get("chunk_hash")
            if content_hash in unique_docs:
                duplicates_removed += 1
            else:
                unique_docs[content_hash] = doc

        final_docs = list(unique_docs.values())

        # Final statistics
        stats.execution_time = time.time() - start_time
        stats.total_chunks = len(final_docs)

        # Get cache statistics
        cache_stats = await self.cache.get_cache_stats()
        stats.cache_hits = cache_stats.get("total_access_count", 0) - len(
            supported_files
        )
        stats.cache_misses = len(supported_files)

        self.logger.info(
            f"Batch processing completed: {stats.processed_files}/{stats.total_files} files, "
            f"{stats.total_chunks} chunks, {duplicates_removed} duplicates removed, "
            f"{stats.execution_time:.2f}s"
        )

        return final_docs, stats

    async def process_directory(
        self,
        directory: Path,
        recursive: bool = True,
        file_pattern: str = "*",
        progress_callback: Optional[ProgressCallback] = None,
        custom_metadata: Optional[Dict[str, Any]] = None,
    ) -> Tuple[List[ProcessedDocument], ProcessingStats]:
        """
        Process all supported files in a directory.

        Args:
            directory: Directory to process
            recursive: Whether to process subdirectories
            file_pattern: File pattern to match
            progress_callback: Optional progress callback
            custom_metadata: Additional metadata for all files

        Returns:
            Tuple of (processed documents, processing statistics)
        """
        if not directory.exists() or not directory.is_dir():
            raise ValueError(f"Invalid directory: {directory}")

        # Find all files
        if recursive:
            file_paths = list(directory.rglob(file_pattern))
        else:
            file_paths = list(directory.glob(file_pattern))

        # Filter to only include files (not directories)
        file_paths = [f for f in file_paths if f.is_file()]

        self.logger.info(f"Found {len(file_paths)} files in {directory}")

        return await self.process_files(
            file_paths,
            progress_callback=progress_callback,
            custom_metadata=custom_metadata,
        )

    def convert_to_documents(
        self, processed_docs: List[ProcessedDocument]
    ) -> List[Document]:
        """Convert ProcessedDocument list to Document list for vector store."""
        documents = []

        for proc_doc in processed_docs:
            doc = Document(
                page_content=proc_doc.content,
                metadata=proc_doc.metadata,
                id=f"{proc_doc.source_file.stem}_{proc_doc.chunk_index}",
            )
            documents.append(doc)

        return documents

    async def get_cache_stats(self) -> Dict[str, Any]:
        """Get processing cache statistics."""
        return await self.cache.get_cache_stats()

    async def clear_cache(self, older_than_days: Optional[int] = None) -> int:
        """Clear processing cache."""
        return await self.cache.clear_cache(older_than_days)


# Convenience functions
async def process_single_file(
    file_path: Path,
    chunk_size: int = 1000,
    overlap: int = 200,
    custom_tags: Optional[Dict[str, str]] = None,
) -> List[Document]:
    """
    Convenience function to process a single file.

    Args:
        file_path: Path to file
        chunk_size: Target chunk size
        overlap: Chunk overlap
        custom_tags: Custom tags to add

    Returns:
        List of Document objects
    """
    processor = DocumentProcessor(
        chunk_size=chunk_size, overlap=overlap, custom_tags=custom_tags
    )

    processed_docs = await processor.process_file(file_path)
    return processor.convert_to_documents(processed_docs)


async def process_directory_simple(
    directory: Path,
    chunk_size: int = 1000,
    overlap: int = 200,
    recursive: bool = True,
    progress_callback: Optional[ProgressCallback] = None,
) -> Tuple[List[Document], ProcessingStats]:
    """
    Convenience function to process a directory.

    Args:
        directory: Directory to process
        chunk_size: Target chunk size
        overlap: Chunk overlap
        recursive: Whether to process subdirectories
        progress_callback: Optional progress callback

    Returns:
        Tuple of (Document list, statistics)
    """
    processor = DocumentProcessor(chunk_size=chunk_size, overlap=overlap)

    processed_docs, stats = await processor.process_directory(
        directory, recursive=recursive, progress_callback=progress_callback
    )

    documents = processor.convert_to_documents(processed_docs)
    return documents, stats


# Export main classes and functions
__all__ = [
    "DocumentProcessor",
    "ProcessedDocument",
    "ProcessingStats",
    "DocumentProcessingError",
    "UnsupportedFormatError",
    "ContentValidationError",
    "DocumentLoader",
    "PDFLoader",
    "TextLoader",
    "DocxLoader",
    "HTMLLoader",
    "TextChunker",
    "MetadataExtractor",
    "TextPreprocessor",
    "ProcessingCache",
    "process_single_file",
    "process_directory_simple",
    "ProgressCallback",
]
